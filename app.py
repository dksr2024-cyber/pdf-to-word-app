from flask import Flask, request, send_file, render_template
import aspose.words as aw
import pytesseract
from PIL import Image
from docx import Document
import os
import re

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# ទិន្នន័យភាសា (បន្ថែមពាក្យសម្រាប់មុខងាររូបភាព)
LANGUAGES = {
    'km': {
        'title': 'បម្លែងឯកសាររហ័ស និងឥតគិតថ្លៃ',
        'pdf_header': 'បម្លែង PDF ទៅជា Word',
        'img_header': 'បម្លែងរូបភាព ទៅជា Word',
        'instruction_pdf': 'សូមជ្រើសរើសឯកសារ PDF',
        'instruction_img': 'សូមជ្រើសរើសរូបភាព (JPG, PNG)',
        'convert_btn': 'បម្លែងឥឡូវនេះ',
        'loading': 'កំពុងបម្លែងឯកសារ សូមរង់ចាំបន្តិច...',
        'support_header': 'គាំទ្រការបង្កើតកម្មវិធីនេះ ☕',
        'support_text': 'សូមអរគុណសម្រាប់ការគាំទ្រ!',
        'account': 'គណនី ABA:',
        'error_empty': 'មិនមានឯកសារត្រូវបានជ្រើសរើសទេ',
        'error_convert': 'មានបញ្ហា៖ '
    },
    'en': {
        'title': 'Fast & Free Converter',
        'pdf_header': 'Convert PDF to Word',
        'img_header': 'Convert Image to Word',
        'instruction_pdf': 'Please select your PDF file',
        'instruction_img': 'Please select your Image (JPG, PNG)',
        'convert_btn': 'Convert Now',
        'loading': 'Converting, please wait...',
        'support_header': 'Support This App ☕',
        'support_text': 'Thank you for your support!',
        'account': 'ABA Account:',
        'error_empty': 'No file selected',
        'error_convert': 'Error: '
    }
}

@app.route('/')
def index():
    lang = request.args.get('lang', 'km')
    if lang not in LANGUAGES: lang = 'km'
    return render_template('index.html', text=LANGUAGES[lang], current_lang=lang)

# --- មុខងារទី១៖ បម្លែង PDF ---
@app.route('/convert-pdf', methods=['POST'])
def convert_pdf():
    lang = request.form.get('lang', 'km')
    text = LANGUAGES.get(lang, LANGUAGES['km'])
    
    file = request.files.get('pdf_file')
    if not file or file.filename == '': return text['error_empty'], 400

    pdf_path = os.path.join(UPLOAD_FOLDER, file.filename)
    docx_filename = file.filename.rsplit('.', 1)[0] + '.docx'
    docx_path = os.path.join(UPLOAD_FOLDER, docx_filename)
    file.save(pdf_path)

    try:
        doc = aw.Document(pdf_path)
        doc.save(docx_path)
        os.remove(pdf_path)
        return send_file(docx_path, as_attachment=True)
    except Exception as e:
        return f"{text['error_convert']} {str(e)}", 500

# --- មុខងារទី២៖ បម្លែង រូបភាព ទៅ Word (ថ្មី) ---
@app.route('/convert-img', methods=['POST'])
def convert_img():
    lang = request.form.get('lang', 'km')
    text = LANGUAGES.get(lang, LANGUAGES['km'])
    
    file = request.files.get('image_file')
    if not file or file.filename == '': return text['error_empty'], 400

    img_path = os.path.join(UPLOAD_FOLDER, file.filename)
    docx_filename = file.filename.rsplit('.', 1)[0] + '.docx'
    docx_path = os.path.join(UPLOAD_FOLDER, docx_filename)
    file.save(img_path)

    try:
        # ១. អានអក្សរពីរូបភាព ដោយប្រាប់ Tesseract ឲ្យស្វែងរកភាសាខ្មែរ និងអង់គ្លេស (khm+eng)
        extracted_text = pytesseract.image_to_string(Image.open(img_path), lang='khm+eng')
        
        # ២. បង្កើតឯកសារ Word ថ្មី រួចសរសេរអក្សរចូល
        doc = Document()
        doc.add_paragraph(extracted_text)
        doc.save(docx_path)
        
        os.remove(img_path)
        return send_file(docx_path, as_attachment=True)
    except Exception as e:
        return f"{text['error_convert']} {str(e)}", 500

if __name__ == '__main__':
    app.run(debug=True, port=5000)
